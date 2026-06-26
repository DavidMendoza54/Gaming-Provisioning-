from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "INTERVIEW_GUIDE.md"
OUTPUT = ROOT / "docs" / "TinyProvisioner_Interview_Guide.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 51)
MUTED = RGBColor(101, 117, 139)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def set_run_font(run, *, name: str = "Calibri", size: float | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    code_style = styles.add_style("CodeBlock", 1)
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code_style.font.size = Pt(9.5)
    code_style.font.color.rgb = RGBColor(17, 24, 39)
    code_style.paragraph_format.left_indent = Inches(0.18)
    code_style.paragraph_format.right_indent = Inches(0.18)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(6)
    code_style.paragraph_format.line_spacing = 1.15

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("TinyProvisioner Interview Guide | ")
    footer_run.font.color.rgb = MUTED
    set_run_font(footer_run, size=9)
    add_page_number(footer)


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("INTERVIEW PREP WORKBOOK")
    run.bold = True
    run.font.color.rgb = BLUE
    set_run_font(run, size=10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run("TinyProvisioner")
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(11, 37, 69)
    set_run_font(title_run, size=30)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    subtitle_run = subtitle.add_run("Compute Provisioning, Docker, Networking, and DevOps Fundamentals")
    subtitle_run.font.color.rgb = MUTED
    set_run_font(subtitle_run, size=13)

    table = doc.add_table(rows=4, cols=2)
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Inches(1.6)
        row.cells[1].width = Inches(4.7)
        for cell in row.cells:
            set_cell_shading(cell, LIGHT_GRAY)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)

    rows = [
        ("Purpose", "Interview-ready explanation of the TinyProvisioner project."),
        ("Audience", "DevOps, security, cloud, and backend internship interviews."),
        ("Format", "Architecture notes, definitions, code map, scenarios, and Q&A."),
        ("Updated", date.today().strftime("%B %d, %Y")),
    ]
    for row, (label, value) in zip(table.rows, rows, strict=True):
        label_run = row.cells[0].paragraphs[0].add_run(label)
        label_run.bold = True
        set_run_font(label_run, size=10.5)
        value_run = row.cells[1].paragraphs[0].add_run(value)
        set_run_font(value_run, size=10.5)

    doc.add_section(WD_SECTION_START.NEW_PAGE)


def add_contents(doc: Document, headings: list[str]) -> None:
    doc.add_heading("Contents", level=1)
    for heading in headings:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(heading)
    doc.add_section(WD_SECTION_START.NEW_PAGE)


def add_markdown_body(doc: Document, source: str) -> None:
    in_code = False
    skipped_title = False

    for raw_line in source.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            continue

        if in_code:
            paragraph = doc.add_paragraph(style="CodeBlock")
            set_paragraph_shading(paragraph, LIGHT_GRAY)
            paragraph.add_run(line if line else " ")
            continue

        if not stripped:
            continue

        if stripped.startswith("# "):
            if not skipped_title:
                skipped_title = True
                continue
            doc.add_heading(stripped[2:], level=1)
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
            continue

        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(0.375)
            paragraph.paragraph_format.first_line_indent = Inches(-0.188)
            paragraph.paragraph_format.space_after = Pt(4)
            add_inline_runs(paragraph, stripped[2:])
            continue

        if stripped[:3].count(".") == 1 and stripped.split(".", 1)[0].isdigit():
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.paragraph_format.left_indent = Inches(0.375)
            paragraph.paragraph_format.first_line_indent = Inches(-0.188)
            paragraph.paragraph_format.space_after = Pt(4)
            add_inline_runs(paragraph, stripped.split(".", 1)[1].strip())
            continue

        paragraph = doc.add_paragraph()
        add_inline_runs(paragraph, stripped)


def add_inline_runs(paragraph, text: str) -> None:
    parts = text.split("`")
    for index, part in enumerate(parts):
        run = paragraph.add_run(part)
        if index % 2 == 1:
            set_run_font(run, name="Consolas", size=9.5)
            run.font.color.rgb = DARK_BLUE
        else:
            set_run_font(run, size=11)


def collect_contents(source: str) -> list[str]:
    headings = []
    for line in source.splitlines():
        if line.startswith("## "):
            headings.append(line[3:].strip())
    return headings


def build() -> None:
    source = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_contents(doc, collect_contents(source))
    add_markdown_body(doc, source)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
